import re
import csv
from docx import Document

INPUT_FILE = 'community.rules'
CSV_FILE = 'snort_rules.csv'
DOCX_FILE = 'snort_rules.docx'

def parse_rule(line):
    """
    Phân tích một dòng rule của Snort, trả về dict với các trường:
    action, protocol, src, src_port, direction, dst, dst_port, msg, sid, rev, raw
    Nếu không phải rule (ví dụ comment hoặc rỗng) trả về None.
    """
    line = line.strip()
    if not line or line.startswith('#'):
        return None
    # tách phần header và options
    m = re.match(r'(\w+)\s+(\w+)\s+(\S+)\s+(\S+)\s+(\-\>|\<\-) \s*(\S+)\s+(\S+)\s*\((.*)\)', line, re.IGNORECASE)
    if not m:
        # không parse được theo mẫu đơn giản
        return {'raw': line}
    action, protocol, src, src_port, direction, dst, dst_port, opts = m.groups()
    # tìm msg, sid, rev
    msg_m = re.search(r'msg\s*:\s*"([^"]+)"', opts)
    sid_m = re.search(r'sid\s*:\s*([0-9]+)', opts)
    rev_m = re.search(r'rev\s*:\s*([0-9]+)', opts)
    msg = msg_m.group(1) if msg_m else ''
    sid = sid_m.group(1) if sid_m else ''
    rev = rev_m.group(1) if rev_m else ''
    return {
        'action': action,
        'protocol': protocol,
        'src': src,
        'src_port': src_port,
        'direction': direction,
        'dst': dst,
        'dst_port': dst_port,
        'msg': msg,
        'sid': sid,
        'rev': rev,
        'raw': line
    }

def export_csv(rules, csv_file):
    fieldnames = ['action','protocol','src','src_port','direction','dst','dst_port','msg','sid','rev','raw']
    with open(csv_file, 'w', newline='', encoding='utf‑8') as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames)
        writer.writeheader()
        for r in rules:
            writer.writerow(r)

def export_docx(rules, docx_file):
    doc = Document()
    doc.add_heading('Snort Rules Export', level=1)
    for r in rules:
        p = doc.add_paragraph()
        p.add_run(f"SID: {r.get('sid','')}, REV: {r.get('rev','')}\n").bold = True
        p.add_run(f"ACTION: {r.get('action','')}, PROTOCOL: {r.get('protocol','')}\n")
        p.add_run(f"SRC: {r.get('src','')}:{r.get('src_port','')} {r.get('direction','')} DST: {r.get('dst','')}:{r.get('dst_port','')}\n")
        p.add_run(f"MSG: {r.get('msg','')}\n")
        p.add_run("Raw rule: \n")
        p.add_run(r.get('raw','') + "\n")
        doc.add_paragraph("\n")
    doc.save(docx_file)

def main():
    rules = []
    with open(INPUT_FILE, encoding='utf‑8', errors='ignore') as f:
        for line in f:
            parsed = parse_rule(line)
            if parsed:
                rules.append(parsed)
    print(f"Parsed {len(rules)} rules.")
    export_csv(rules, CSV_FILE)
    print(f"Exported CSV to {CSV_FILE}")
    export_docx(rules, DOCX_FILE)
    print(f"Exported DOCX to {DOCX_FILE}")

if __name__ == '__main__':
    main()
